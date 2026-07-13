import os
import base64
import time
import json
from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn
from PIL import Image
import io

app = Flask(__name__, static_folder='static', template_folder='templates')

DB_FILE = 'questions.json'
PDF_DIR = 'pdfs'
CROPPED_DIR = os.path.join('static', 'cropped')

# 確保目錄存在
os.makedirs(PDF_DIR, exist_ok=True)
os.makedirs(CROPPED_DIR, exist_ok=True)

if not os.path.exists(DB_FILE):
    with open(DB_FILE, 'w', encoding='utf-8') as f:
        json.dump([], f)

def read_db():
    try:
        with open(DB_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []

def write_db(data):
    with open(DB_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

@app.route('/')
def index():
    return render_template('index.html')

import shutil

@app.route('/diff/<version>', defaults={'year_str': '112'})
@app.route('/diff/<version>/<year_str>')
def diff_version(version, year_str):
    year_map = {
        '108': '2019',
        '109': '2020',
        '110': '2021',
        '111': '2022',
        '112': '2023',
        '113': '2024',
        '114': '2025'
    }
    
    bc_year = year_map.get(year_str, '2023')
    user_src_dir = r"D:\OneDrive\01篩選題目\切圖測試結果"
    user_dest_dir = os.path.join('static', f'user_{year_str}')
    os.makedirs(user_dest_dir, exist_ok=True)
    
    questions = []
    for q in range(1, 26):
        user_filename = f"國{bc_year} ({q}).png"
        user_src_path = os.path.join(user_src_dir, user_filename)
        user_dest_filename = f"user_q{q}.png"
        user_dest_path = os.path.join(user_dest_dir, user_dest_filename)
        
        if os.path.exists(user_src_path):
            try:
                shutil.copy2(user_src_path, user_dest_path)
            except Exception:
                pass
            user_url = f"/static/user_{year_str}/{user_dest_filename}"
        else:
            user_url = ""
            
        test_filename = f"q_{year_str}_國語_一年級_{q}.png"
        test_rel_path = f"static/cropped_{version}/{test_filename}"
        test_exists = os.path.exists(test_rel_path)
        if not test_exists:
            test_filename = f"q_{year_str}_國語_一年級_{q}_v143.png"
            test_rel_path = f"static/cropped_{version}/{test_filename}"
            test_exists = os.path.exists(test_rel_path)
        test_url = f"/static/cropped_{version}/{test_filename}?t={int(os.path.getmtime(test_rel_path))}" if test_exists else ""
        
        reading_filename = f"q_{year_str}_國語_一年級_reading_2.png"
        reading_rel_path = f"static/cropped_{version}/{reading_filename}"
        reading_exists = os.path.exists(reading_rel_path)
        reading_url = f"/static/cropped_{version}/{reading_filename}?t={int(os.path.getmtime(reading_rel_path))}" if reading_exists else ""
        
        questions.append({
            'q': q,
            'user_url': user_url,
            'test_url': test_url,
            'test_exists': test_exists,
            'reading_url': reading_url,
            'reading_exists': reading_exists
        })
        
    return render_template('diff.html', version=version, year=year_str, questions=questions)




@app.route('/diagnostic')
def diagnostic():
    return render_template('diagnostic.html')

@app.route('/api/pdfs', methods=['GET'])
def list_pdfs():
    """列出 pdfs/ 目錄下的所有 PDF 檔案"""
    files = [f for f in os.listdir(PDF_DIR) if f.lower().endswith('.pdf')]
    return jsonify(files)

@app.route('/api/pdfs/upload', methods=['POST'])
def upload_pdf():
    """上傳 PDF 檔案至 pdfs/ 目錄"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '沒有檔案部分'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未選擇檔案'}), 400
        if file and file.filename.lower().endswith('.pdf'):
            filepath = os.path.join(PDF_DIR, file.filename)
            file.save(filepath)
            return jsonify({'status': 'success', 'filename': file.filename})
        return jsonify({'error': '只支援 PDF 格式'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/questions', methods=['GET'])
def get_questions():
    """取得所有題目資料庫 (並附帶時間戳以防瀏覽器快取舊圖片)"""
    db = read_db()
    for q in db:
        if 'imagePaths' in q:
            q['imagePaths'] = [f"{path}?t={int(os.path.getmtime(path.lstrip('/')))}" if os.path.exists(path.lstrip('/')) else f"{path}?t={int(time.time())}" for path in q['imagePaths']]
        if 'imagePath' in q:
            path = q['imagePath']
            q['imagePath'] = f"{path}?t={int(os.path.getmtime(path.lstrip('/')))}" if os.path.exists(path.lstrip('/')) else f"{path}?t={int(time.time())}"
    return jsonify(db)

@app.route('/api/questions/save', methods=['POST'])
def save_question():
    """儲存裁切的題目圖片與標籤"""
    try:
        data = request.json
        if not data or 'image' not in data:
            return jsonify({'error': '缺少圖片資料'}), 400
        
        # 解析 Base64 圖片
        image_data = data['image']
        if ',' in image_data:
            image_data = image_data.split(',')[1]
        
        # 產生唯一檔名
        timestamp = int(time.time() * 1000)
        filename = f"q_{timestamp}.png"
        filepath = os.path.join(CROPPED_DIR, filename)
        
        # 寫入檔案
        with open(filepath, 'wb') as f:
            f.write(base64.b64decode(image_data))
        
        # 組裝 Metadata
        new_q = {
            'id': f"q_{timestamp}",
            'imagePath': f"/static/cropped/{filename}",
            'year': data.get('year', '').strip(),
            'subject': data.get('subject', '').strip(),
            'grade': data.get('grade', '').strip(),
            'type': data.get('type', '').strip(),
            'tags': [t.strip() for t in data.get('tags', []) if t.strip()],
            'createdAt': time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
        }
        
        # 存入資料庫
        db = read_db()
        db.append(new_q)
        write_db(db)
        
        return jsonify({'status': 'success', 'question': new_q})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/questions/delete', methods=['POST'])
def delete_question():
    """刪除特定題目"""
    try:
        data = request.json
        if not data or 'id' not in data:
            return jsonify({'error': '缺少題目 ID'}), 400
        
        q_id = data['id']
        db = read_db()
        
        # 尋找題目
        target = None
        for q in db:
            if q['id'] == q_id:
                target = q
                break
                
        if not target:
            return jsonify({'error': '找不到該題目'}), 404
            
        # 刪除圖片檔案
        # 將 /static/cropped/... 轉為實體路徑
        rel_path = target['imagePath'].lstrip('/')
        if os.path.exists(rel_path):
            os.remove(rel_path)
            
        # 從資料庫移除
        db = [q for q in db if q['id'] != q_id]
        write_db(db)
        
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export-docx', methods=['POST'])
def export_docx():
    """將選取的題目打包成 Word 檔下載"""
    try:
        data = request.json
        if not data or 'questionIds' not in data:
            return jsonify({'error': '缺少題目清單'}), 400
            
        title = data.get('title', '考古題練習卷').strip()
        subtitle = data.get('subtitle', '').strip()
        q_ids = data['questionIds']
        
        # 取得所有選取的題目資料並保持前端的排序
        db = read_db()
        q_dict = {q['id']: q for q in db}
        selected_questions = []
        for q_id in q_ids:
            if q_id in q_dict:
                selected_questions.append(q_dict[q_id])
                
        if not selected_questions:
            return jsonify({'error': '選取的題目清單為空'}), 400
            
        # 建立 Word 文件
        doc = Document()
        
        # 設定邊距與 B4 Portrait 紙張大小 (25.7 x 36.4 公分 = 10.12 x 14.33 英吋)
        section1 = doc.sections[0]
        section1.orientation = WD_ORIENT.PORTRAIT
        section1.page_width = Inches(10.12)
        section1.page_height = Inches(14.33)
        section1.top_margin = Inches(0.6)
        section1.bottom_margin = Inches(0.6)
        section1.left_margin = Inches(0.6)
        section1.right_margin = Inches(0.6)
            
        # 1. 考卷標題 (標楷體)
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title)
        run_title.font.name = '標楷體'
        
        # 動態調整字體大小以確保只呈現在同一行
        title_len = len(title)
        if title_len > 40:
            title_size = 14
        elif title_len > 30:
            title_size = 17
        else:
            title_size = 22
            
        run_title.font.size = Pt(title_size)
        run_title.bold = True
        rPr = run_title._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '標楷體')
        
        # 2. 副標題 (標楷體)
        if subtitle:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(subtitle)
            run_sub.font.name = '標楷體'
            run_sub.font.size = Pt(12)
            rPr = run_sub._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '標楷體')
            
        # 3. 學生資訊欄位表格 (B4 高度放大，標楷體 16 號字，方便書寫)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        
        # 設定行高為 0.45 英吋
        row = table.rows[0]
        row.height = Inches(0.45)
        
        hdr_cells = row.cells
        hdr_cells[0].text = ' 班級：'
        hdr_cells[1].text = ' 座號：'
        hdr_cells[2].text = ' 姓名：'
        hdr_cells[3].text = ' 得分：'
        
        # 微調字型與單格內距 (標楷體 16 號字)
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = '標楷體'
                run.font.size = Pt(16)
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), '標楷體')
                         
        # 空一行
        doc.add_paragraph()
        
        # 建立雙欄排版 Section 2 (接續在同一頁，B4 直式)
        section2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
        section2.orientation = WD_ORIENT.PORTRAIT
        section2.page_width = Inches(10.12)
        section2.page_height = Inches(14.33)
        section2.top_margin = Inches(0.6)
        section2.bottom_margin = Inches(0.6)
        section2.left_margin = Inches(0.6)
        section2.right_margin = Inches(0.6)
        
        # 低階 XML 設定雙欄、間距與垂直分隔線 (sep="1")
        sectPr = section2._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            from docx.oxml import OxmlElement
            cols_el = OxmlElement('w:cols')
            sectPr.append(cols_el)
        else:
            cols_el = cols[0]
        cols_el.set(qn('w:num'), '2')
        cols_el.set(qn('w:space'), '576')
        cols_el.set(qn('w:sep'), '1')
        
        # B4 直式雙欄下，單欄最大可列印寬度約為 4.26 英吋
        MAX_COL_WIDTH = 4.26
        
        # 4. 依序寫入題目 (直接插入截圖，不附加編號，維持圖片整齊)
        for q in selected_questions:
            img_paths = q.get('imagePaths', [])
            if not img_paths and 'imagePath' in q:
                img_paths = [q['imagePath']]
                
            for img_path_rel in img_paths:
                img_path = img_path_rel.lstrip('/')
                if os.path.exists(img_path):
                    p_q = doc.add_paragraph()
                    p_q.paragraph_format.space_after = Pt(8)
                    try:
                        with Image.open(img_path) as img:
                            w_px, h_px = img.size
                        # 計算寬度，自動等比例縮小以適應 B4 直式單欄
                        dpi = 130
                        calc_w = w_px / dpi
                        final_w = min(calc_w, MAX_COL_WIDTH)
                        
                        p_q.add_run().add_picture(img_path, width=Inches(final_w))
                    except Exception as img_err:
                        p_q.add_run(f"[無法載入圖片: {img_err}]")
                else:
                    p_q = doc.add_paragraph()
                    p_q.add_run("[找不到題目圖片]")
            
        # 儲存至記憶體 buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 回傳 Word 檔案流
        filename = f"exam_{int(time.time())}.docx"
        return send_file(
            file_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # 啟動在本機的 5000 port
    app.run(host='127.0.0.1', port=5000, debug=True)
